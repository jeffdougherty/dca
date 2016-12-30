from docx import *
from helperfunctions import connect_to_db, deepcopy
from tkintertable import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocsPrinter(Frame):
    #Takes the scenario data provided by ScenariosScreen, retrieves additional data, and prepares for printing.

    def __init__(self, parent, data):
        self.parent = parent
        self.master = self.parent.master
        #Have the class automatically find the other data it needs to print as soon as it starts up.  Store what you need in other methods, throw the rest out as soon as init closes
        self.scenario_key = data['Scenario Key']
        del data['Scenario Key'] #Removing metadata so the user-facing data can be stored
        self.scenario_data = data
        cursor = connect_to_db()
        self.sides_keys = []
        self.sides_data = {}
        self.formation_data = {}
        #Get the sides data first
        cursor.execute("""SELECT * FROM 'Scenario Side' WHERE [Scenario Key]=?""", (self.scenario_key,))
        column_titles = [[description[0]] for description in cursor.description]
        for side in cursor.fetchall():
            working_copy = self.parse_data_for_print(side, column_titles)
            del working_copy['Scenario Key']
            this_side_key = working_copy['Side Key']
            self.sides_keys.append(this_side_key)
            del working_copy['Side Key']
            self.sides_data[this_side_key] = working_copy
        #Now we get the formations data
        for this_side_key in self.sides_keys:
            side_formations = {}
            cursor.execute("""SELECT * FROM 'Scenario Ship Formation' WHERE [Scenario Key]=? AND [Scenario Side]=?""", (self.scenario_key, this_side_key,))
            column_titles = [[description[0]] for description in cursor.description]
            for formation in cursor.fetchall():
                working_copy = self.parse_data_for_print(formation, column_titles)
                del working_copy['Scenario Key']
                del working_copy['Scenario Side']
                this_formation_key = working_copy['Formation ID']
                del working_copy['Formation ID']
                side_formations[this_formation_key] = working_copy
                self.get_formation_ships(working_copy, this_side_key, this_formation_key)
            self.formation_data[this_side_key] = side_formations
        #Now have the user select which documents need to be created
        #Then have functions to generate them
        #GM's Brief
        #Side Briefs
        Frame.__init__(self, parent, background='white')
        self.pack(fill=BOTH, expand=1)
        self.adjust_geometry()
        self.checkbox_frame = Frame(self, width = self.width / 2, height = self.height)
        self.checkbox_frame.pack(side='left')
        self.button_frame = Frame(self, width=self.width/2, height = self.height)
        self.button_frame.pack(side='left')
        self.gm_brief_var = IntVar()
        self.gm_brief_box = Checkbutton(self.checkbox_frame, text="GM Brief", variable=self.gm_brief_var)
        self.gm_brief_box.pack(side='left')
        self.side_brief_vars = {}
        self.side_brief_boxes = {}
        for key in self.sides_keys:
            self.side_brief_vars[key] = IntVar()
            this_side_name = self.sides_data[key]['Side Name']
            self.side_brief_boxes[key] = Checkbutton(self.checkbox_frame, text=this_side_name, variable=self.side_brief_vars[key])
            self.side_brief_boxes[key].pack(side='left')
        OKButton = Button(self.button_frame, text = 'OK', command=lambda: self.print_var_states())
        OKButton.pack(side='top')
        CancelButton = Button(self.button_frame, text='Cancel', command=lambda: self.close_window())
        CancelButton.pack(side='top')

    def adjust_geometry(self):
        self.width = self.parent.winfo_width()
        self.height = self.parent.winfo_toplevel().winfo_height()
        self.parent.geometry = self.winfo_toplevel().winfo_geometry()

    def parse_data_for_print(self, data_tuple, column_list):
        data_dict = {}
        assert len(data_tuple) == len(column_list), "Title list and data list do not match!"
        for index in range(len(column_list)):
            this_key = column_list[index][0]
            data_dict[this_key] = data_tuple[index]
        return data_dict

    def get_formation_ships(self, formation_data_dict, this_side_key, this_formation_key):
        formation_ships = {}
        cursor = connect_to_db()
        cursor.execute("""SELECT * FROM 'Scenario Ship Formation Ship' WHERE [Scenario Key]=? AND [Scenario Side]=? AND [Formation ID]=?""",(self.scenario_key, this_side_key, this_formation_key,))
        column_titles = [[description[0]] for description in cursor.description]
        ship_column_titles = None
        ship_columns_to_retrieve = ['Ship Type', 'Class', 'Class Variant', 'Damage Pts', 'Speed']
        for ship in cursor.fetchall():
            working_copy = self.parse_data_for_print(ship, column_titles)
            del working_copy['Scenario Key']
            del working_copy['Scenario Side']
            del working_copy['Formation ID']
            this_ship_key = working_copy['Formation Ship Key']
            del working_copy['Formation Ship Key']
            this_annex_a_key = working_copy['Annex A Key']
            del working_copy['Annex A Key']
            cursor.execute("""SELECT * FROM 'Ship' WHERE [Ship Key]=?""", (this_annex_a_key,))
            if ship_column_titles == None:
                ship_column_titles = [[description[0]] for description in cursor.description]
            ship_data = self.parse_data_for_print(cursor.fetchone(), ship_column_titles) #Can use fetchone() here since we know we'll be returning only one record
            for entry in ship_columns_to_retrieve:
                working_copy[entry] = ship_data[entry]
            formation_ships[this_ship_key] = working_copy
        formation_data_dict['Ships'] = formation_ships
        return formation_data_dict

    def center_window(self):
        w, h = self.width, self.height
        self.sw = self.parent.winfo_screenwidth()
        self.sh = self.parent.winfo_screenheight()
        x = (self.sw - self.width) / 2
        y = (self.sh - self.height) / 2
        self.parent.geometry('%dx%d+%d+%d' % (w, h, x, y))

    def close_window(self):
        self.destroy()
        self.parent.destroy()

    def print_var_states(self):
        print "GM Brief: " + str(self.gm_brief_var.get())
        for this_key in self.side_brief_boxes.keys():
            this_side_name = self.side_brief_boxes[this_key]['text']
            print this_side_name + " Brief: " + str(self.side_brief_vars[this_key].get())
        self.get_docs_to_make()

    def get_docs_to_make(self):
        if self.gm_brief_var.get() == True:
            self.make_gm_brief()
        for this_key in self.side_brief_boxes.keys():
            if self.side_brief_vars[this_key].get() == True:
                self.make_side_brief(this_key)
        self.close_window()

    def make_gm_brief(self):
        GM_Brief = Document()
        self.add_title_intro_and_basic_info(GM_Brief, True)
        for side_key in self.sides_keys:
            self.add_side_info(GM_Brief, side_key)
        GM_Brief.save('~/GM Brief.docx')

    def make_side_brief(self, this_side_key):
        Player_Brief = Document()
        this_side_name = self.sides_data[this_side_key]['Side Name']
        self.add_title_intro_and_basic_info(Player_Brief, False)
        self.add_side_info(Player_Brief, this_side_key)
        file_name = this_side_name + " Brief.docx"
        Player_Brief.save(file_name)

    def add_title_intro_and_basic_info(self, this_doc, show_end_time=False):
        title = this_doc.add_heading(self.scenario_data['Scenario Title'], 0)
        title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p = this_doc.add_paragraph("")
        p.add_run("Introduction: ").bold = True
        p.add_run(self.scenario_data['Intro Text'])
        p = this_doc.add_paragraph("")
        p.add_run("Location: ").bold = True
        location_string = str(self.scenario_data['Location']) + ", " + str(self.scenario_data['Start Date Time'])
        if show_end_time:
            location_string += " -- " + self.scenario_data['End Date Time']
        p.add_run(location_string)
        environment_variables = ['Sea State', 'Visibility Percent', 'Wind Data', 'Obscuration', 'Sonar Modifier']
        air_environment_variables = ['Ceiling', 'Cloud Coverage Percent']
        wind_data_formatted = str(self.scenario_data['Wind Velocity Base']) + " kts from " + str(self.scenario_data['Wind Direction'])
        if self.scenario_data['Wind Velocity Gusts'] != None:
            wind_data_formatted += ', gusts to ' + str(self.scenario_data['Wind Velocity Gusts'])
        environment_string = ""
        air_environment_string = ""
        for this_var in environment_variables:
            if this_var == 'Wind Data':
                if environment_string != "":
                    environment_string += ", "
                environment_string += "Wind " + wind_data_formatted
            elif self.scenario_data[this_var] != None:
                if environment_string != "":
                    environment_string += ", "
                if "Percent" in this_var:
                    this_var_print = this_var.replace("Percent", "") #Want "Visibility", not "Visibility Percent"
                    environment_string += this_var_print + " " +str(self.scenario_data[this_var]) + "%"
                else:
                    environment_string += this_var + " " + str(self.scenario_data[this_var])
        for this_var in air_environment_variables:
            if self.scenario_data[this_var] != None:
                if air_environment_string != "":
                    air_environment_string += ", "
                if "Percent" in this_var:
                    this_var_print = this_var.replace("Percent", "")  # Want "Visibility", not "Visibility Percent"
                    air_environment_string += this_var_print + " " + str(self.scenario_data[this_var]) + "%"
                else:
                    air_environment_string += this_var + " " + str(self.scenario_data[this_var])
        p = this_doc.add_paragraph("")
        p.add_run("Environment: ").bold = True
        p.add_run(environment_string)
        if len(air_environment_string) > 0:
            p = this_doc.add_paragraph(air_environment_string)

    def add_side_info(self, this_doc, this_side_key):
        this_side_data = deepcopy(self.sides_data[this_side_key])
        this_side_name = this_side_data['Side Name']
        del this_side_data['Side Name']
        self.add_side_info_element(this_doc, this_side_data, "Side Situation Report", this_side_name)
        setup_string, additional_orders_string = self.add_side_forces(this_doc, this_side_key)
        if len(additional_orders_string) > 0:
            orders_graph = self.add_side_info_element(this_doc, this_side_data, "Side Orders", this_side_name, return_p=True)
            orders_graph.add_run(additional_orders_string)
        else:
            self.add_side_info_element(this_doc, this_side_data, "Side Orders", this_side_name)
        self.add_side_info_element(this_doc, this_side_data, "Side Victory Conditions", this_side_name)
        """

        for this_key in this_side_data.keys():
            this_label = str(this_side_name) + " " + this_key.replace("Side ", "") + ": "
            this_data = str(this_side_data[this_key])
            print this_label + this_data
            p = this_doc.add_paragraph("")
            p.add_run(this_label).bold = True
            p.add_run(this_data)
        """
        p = this_doc.add_paragraph("")
        p.add_run("Setup: ").bold = True
        p.add_run(setup_string)

    def add_side_info_element(self, this_doc, side_data, element_key, side_name, return_p = False):
        this_label = str(side_name) + " " + element_key.replace("Side ", "") + ": "
        this_data = str(side_data[element_key])
        p = this_doc.add_paragraph("")
        p.add_run(this_label).bold = True
        p.add_run(this_data)
        if return_p:
            return p

    def add_side_forces(self, this_doc, this_side_key):
        this_formation_data = deepcopy(self.formation_data[this_side_key])
        forces_p = this_doc.add_paragraph("")
        setup_string = ""
        formation_orders_string = ""
        for formation in this_formation_data.keys():
            #Write out the formation itself
            this_data = this_formation_data[formation]
            if this_data['Formation Name'] != None and len(this_data['Formation Name']) > 0:
                forces_p.add_run(this_data['Formation Name'] + ": ")
            is_first_ship = True
            for ship in this_data['Ships']:
                this_ship = this_data['Ships'][ship]
                if not is_first_ship:
                    forces_p.add_run(", ")
                is_first_ship = False
                forces_p.add_run(this_ship['Ship Name']).italic = True
                forces_p.add_run(" (" + str(this_ship['Class']) + " class " + str(this_ship['Ship Type']))
                if this_ship['Class Variant'] != None and len(this_ship['Class Variant']) > 0:
                    forces_p.add_run(", " + str(this_ship['Class Variant']))
                forces_p.add_run(")")
            forces_p.add_run(". ")
            #Save the setup data for later use
            this_formation_setup = this_data['Formation Name'] + ": " + this_data["Formation Starting Location"] + ". Course " + str(this_data["Formation Course"]) + ", speed " + str(this_data["Formation Speed"]) + " knots. "
            orders_lowercase = "".join(c.lower() if i == 0 else c for i, c in enumerate(this_data['Formation Orders'])) #Dropping an initial capital letter is not the easiest thing to do in Python
            this_formation_orders = this_data['Formation Name'] + " will " + orders_lowercase
            setup_string += this_formation_setup
            formation_orders_string += this_formation_orders
        return setup_string, formation_orders_string